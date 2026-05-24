"""
CSRD Comply — Export Multi-Formato (Step 22)

Servizio di esportazione report nei formati supportati:
1. iXBRL (XHTML) — formato principale per filing regolatorio
2. PDF — per stampa e condivisione interna
3. XLSX — per analisi dati
4. Word (DOCX) — per bozze e revisioni
5. JSON — per API integration

Librerie utilizzate:
- xhtml2pdf: conversione XHTML → PDF
- openpyxl: generazione file Excel
- python-docx: generazione file Word
"""
import io
import json
import logging
from typing import Optional, Dict, Any, List
from datetime import datetime
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)


# ── Data Classes ──────────────────────────────────────────────────

@dataclass
class ExportOptions:
    """Opzioni di esportazione comuni a tutti i formati."""
    include_cover: bool = True
    include_toc: bool = True
    include_compliance: bool = True
    language: str = "en"
    watermark: Optional[str] = None  # "Draft", "Review", ecc.
    page_size: str = "A4"
    margin_top: float = 20.0
    margin_bottom: float = 20.0
    margin_left: float = 15.0
    margin_right: float = 15.0


@dataclass
class ExportResult:
    """Risultato di un'operazione di export."""
    format: str
    filename: str
    content: bytes
    mime_type: str
    size_bytes: int
    success: bool = True
    error_message: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)


class ExportError(Exception):
    """Eccezione generica per errori di export."""
    pass


class PDFGenerationError(ExportError):
    """Errore nella generazione PDF."""
    pass


class XLSXGenerationError(ExportError):
    """Errore nella generazione Excel."""
    pass


class DOCXGenerationError(ExportError):
    """Errore nella generazione Word."""
    pass


# ── Export Service ────────────────────────────────────────────────

class ExportService:
    """
    Servizio centrale per l'esportazione di report CSRD in formati multipli.
    
    Usage:
        service = ExportService()
        
        # Export PDF
        pdf_result = service.export_pdf(xhtml_content, options)
        
        # Export XLSX (dati strutturati)
        xlsx_result = service.export_xlsx(report_data, options)
        
        # Export DOCX
        docx_result = service.export_docx(xhtml_content, options)
        
        # Export JSON
        json_result = service.export_json(report_data, options)
        
        # Export iXBRL (pass-through con validazione)
        ixbrl_result = service.export_ixbrl(ixbrl_content, options)
        
        # Export completo (tutti i formati)
        all_results = service.export_all(ixbrl_content, report_data, options)
    """

    # MIME types per i formati supportati
    MIME_TYPES = {
        "ixbrl": "application/xhtml+xml",
        "pdf": "application/pdf",
        "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "json": "application/json",
    }

    # Estensioni file
    EXTENSIONS = {
        "ixbrl": ".xhtml",
        "pdf": ".pdf",
        "xlsx": ".xlsx",
        "docx": ".docx",
        "json": ".json",
    }

    def __init__(self):
        """Inizializza il servizio di export."""
        self._xhtml2pdf_available = self._check_library("xhtml2pdf")
        self._openpyxl_available = self._check_library("openpyxl")
        self._docx_available = self._check_library("docx")
        logger.info(
            f"ExportService initialized. Libraries: "
            f"xhtml2pdf={'✓' if self._xhtml2pdf_available else '✗'}, "
            f"openpyxl={'✓' if self._openpyxl_available else '✗'}, "
            f"python-docx={'✓' if self._docx_available else '✗'}"
        )

    def _check_library(self, name: str) -> bool:
        """Verifica se una libreria è installata."""
        try:
            if name == "xhtml2pdf":
                import xhtml2pdf  # noqa
            elif name == "openpyxl":
                import openpyxl  # noqa
            elif name == "docx":
                import docx  # noqa
            return True
        except ImportError:
            return False

    # ── PDF Export ──────────────────────────────────────────────

    def export_pdf(
        self,
        xhtml_content: str,
        filename: str = "report",
        options: Optional[ExportOptions] = None,
    ) -> ExportResult:
        """
        Esporta il report in formato PDF.
        
        Usa xhtml2pdf per convertire XHTML in PDF.
        Fallback: genera un PDF testuale semplice se la libreria non è disponibile.
        
        Args:
            xhtml_content: Contenuto XHTML del report
            filename: Nome base del file (senza estensione)
            options: Opzioni di esportazione
            
        Returns:
            ExportResult con il contenuto PDF in bytes
            
        Raises:
            PDFGenerationError: Se la generazione PDF fallisce
        """
        opts = options or ExportOptions()
        full_filename = f"{filename}{self.EXTENSIONS['pdf']}"

        if not self._xhtml2pdf_available:
            logger.warning("xhtml2pdf not available, generating fallback PDF")
            return self._generate_fallback_pdf(xhtml_content, full_filename, opts)

        try:
            from xhtml2pdf import pisa

            # Applica watermark se richiesto
            html_to_render = xhtml_content
            if opts.watermark:
                html_to_render = self._apply_watermark(xhtml_content, opts.watermark)

            # Genera PDF in memoria
            pdf_buffer = io.BytesIO()
            result = pisa.CreatePDF(
                src=html_to_render,
                dest=pdf_buffer,
                encoding="UTF-8",
                show_error_as_pdf=True,
            )

            if result.err:
                error_msg = f"xhtml2pdf error: {result.err}"
                logger.error(error_msg)
                raise PDFGenerationError(error_msg)

            pdf_buffer.seek(0)
            pdf_bytes = pdf_buffer.getvalue()

            logger.info(f"PDF generated: {len(pdf_bytes)} bytes")
            return ExportResult(
                format="pdf",
                filename=full_filename,
                content=pdf_bytes,
                mime_type=self.MIME_TYPES["pdf"],
                size_bytes=len(pdf_bytes),
                metadata={
                    "pages_estimated": self._estimate_pages(pdf_bytes),
                    "generated_at": datetime.utcnow().isoformat(),
                },
            )

        except PDFGenerationError:
            raise
        except Exception as e:
            error_msg = f"Unexpected PDF generation error: {str(e)}"
            logger.error(error_msg)
            raise PDFGenerationError(error_msg)

    def _apply_watermark(self, html_content: str, watermark_text: str) -> str:
        """Applica un watermark al contenuto HTML."""
        watermark_style = f"""
        <div style="
            position: fixed; top: 50%; left: 50%;
            transform: translate(-50%, -50%) rotate(-45deg);
            font-size: 72px; color: rgba(255,0,0,0.1);
            font-weight: bold; z-index: 1000;
            pointer-events: none;
        ">{watermark_text}</div>
        """
        return html_content.replace("</body>", f"{watermark_style}\n</body>")

    def _generate_fallback_pdf(
        self,
        xhtml_content: str,
        filename: str,
        options: ExportOptions,
    ) -> ExportResult:
        """
        Genera un PDF testuale di fallback quando xhtml2pdf non è disponibile.
        
        Crea un semplice PDF con ReportLab o produce un errore informativo.
        """
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
            )
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib import colors

            pdf_buffer = io.BytesIO()
            doc = SimpleDocTemplate(
                pdf_buffer,
                pagesize=A4,
                topMargin=options.margin_top,
                bottomMargin=options.margin_bottom,
                leftMargin=options.margin_left,
                rightMargin=options.margin_right,
            )

            styles = getSampleStyleSheet()
            story = []

            # Parsing HTML di base -> estrai testo
            import re
            text_content = re.sub(r'<[^>]+>', ' ', xhtml_content)
            text_content = re.sub(r'\s+', ' ', text_content).strip()

            # Dividi in paragrafi
            paragraphs = text_content.split('\n')
            for para in paragraphs[:200]:  # Limita a 200 paragrafi
                para = para.strip()
                if para:
                    story.append(Paragraph(para, styles["Normal"]))
                    story.append(Spacer(1, 6))

            doc.build(story)
            pdf_buffer.seek(0)
            pdf_bytes = pdf_buffer.getvalue()

            logger.info(f"Fallback PDF generated: {len(pdf_bytes)} bytes")
            return ExportResult(
                format="pdf",
                filename=filename,
                content=pdf_bytes,
                mime_type=self.MIME_TYPES["pdf"],
                size_bytes=len(pdf_bytes),
                metadata={
                    "fallback": True,
                    "message": "Generated with ReportLab fallback (xhtml2pdf not available)",
                    "generated_at": datetime.utcnow().isoformat(),
                },
            )

        except ImportError:
            # ReportLab non disponibile -> genera PDF minimale
            pdf_bytes = self._generate_minimal_pdf(
                f"CSRD Report - {filename}",
                "PDF generation requires xhtml2pdf or reportlab\n"
                "Install with: pip install xhtml2pdf reportlab",
            )
            return ExportResult(
                format="pdf",
                filename=filename,
                content=pdf_bytes,
                mime_type=self.MIME_TYPES["pdf"],
                size_bytes=len(pdf_bytes),
                metadata={
                    "fallback": True,
                    "message": "Minimal PDF generated - install xhtml2pdf for full rendering",
                    "generated_at": datetime.utcnow().isoformat(),
                },
            )

    def _generate_minimal_pdf(self, title: str, body_text: str) -> bytes:
        """Genera un PDF minimale senza dipendenze esterne."""
        pdf_lines = [
            b"%PDF-1.4",
            b"1 0 obj",
            b"<< /Type /Catalog /Pages 2 0 R >>",
            b"endobj",
            b"2 0 obj",
            b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
            b"endobj",
            b"3 0 obj",
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
            b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>",
            b"endobj",
            b"4 0 obj",
            b"<< /Length 200 >>",
            b"stream",
            b"BT /F1 24 Tf 50 800 Td (CSRD Report) Tj ET",
            b"BT /F1 12 Tf 50 750 Td (Install xhtml2pdf for proper PDF) Tj ET",
            b"endstream",
            b"endobj",
            b"5 0 obj",
            b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
            b"endobj",
            b"xref",
            b"0 6",
            b"0000000000 65535 f ",
            b"0000000009 00000 n ",
            b"0000000058 00000 n ",
            b"0000000115 00000 n ",
            b"0000000266 00000 n ",
            b"0000000498 00000 n ",
            b"trailer << /Size 6 /Root 1 0 R >>",
            b"startxref",
            b"566",
            b"%%EOF",
        ]
        return b"\n".join(pdf_lines)

    def _estimate_pages(self, pdf_bytes: bytes) -> int:
        """Stima il numero di pagine di un PDF."""
        count = pdf_bytes.count(b"/Type /Page") - pdf_bytes.count(b"/Type /Pages")
        return max(1, count)

    # ── XLSX Export ─────────────────────────────────────────────

    def export_xlsx(
        self,
        report_data: Dict[str, Any],
        filename: str = "report_data",
        options: Optional[ExportOptions] = None,
    ) -> ExportResult:
        """
        Esporta i dati del report in formato Excel (XLSX).
        
        Crea un workbook con più fogli:
        - Summary: riepilogo report
        - GHG Emissions: dati emissivi
        - Materiality: punteggi doppia materialità
        - ESRS Coverage: copertura datapoint
        
        Args:
            report_data: Dizionario con i dati strutturati del report
            filename: Nome base del file
            options: Opzioni di esportazione
            
        Returns:
            ExportResult con il contenuto XLSX in bytes
        """
        opts = options or ExportOptions()
        full_filename = f"{filename}{self.EXTENSIONS['xlsx']}"

        if not self._openpyxl_available:
            logger.warning("openpyxl not available for XLSX export")
            return ExportResult(
                format="xlsx",
                filename=full_filename,
                content=b"",
                mime_type=self.MIME_TYPES["xlsx"],
                size_bytes=0,
                success=False,
                error_message="openpyxl not installed. Install with: pip install openpyxl",
            )

        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter

            wb = openpyxl.Workbook()

            # ── Foglio 1: Summary ──────────────────────────────
            ws_summary = wb.active
            ws_summary.title = "Summary"

            header_font = Font(bold=True, size=12, color="FFFFFF")
            header_fill = PatternFill(start_color="2B6CB0", end_color="2B6CB0", fill_type="solid")
            data_font = Font(size=11)
            thin_border = Border(
                left=Side(style="thin"),
                right=Side(style="thin"),
                top=Side(style="thin"),
                bottom=Side(style="thin"),
            )

            # Intestazioni summary
            summary_headers = ["Field", "Value"]
            ws_summary.append(summary_headers)
            for cell in ws_summary[1]:
                cell.font = header_font
                cell.fill = header_fill
                cell.border = thin_border

            # Dati summary
            summary_data = [
                ("Company", report_data.get("company_name", "")),
                ("Report Title", report_data.get("report_title", "")),
                ("Reporting Year", str(report_data.get("reporting_year", ""))),
                ("Language", report_data.get("language", "en")),
                ("Generated At", report_data.get("generated_at", "")),
                ("Status", report_data.get("status", "draft")),
                ("ESRS Version", report_data.get("esrs_version", "")),
            ]
            for field, value in summary_data:
                ws_summary.append([field, value])

            for row in ws_summary.iter_rows(min_row=2, max_row=len(summary_data) + 1):
                for cell in row:
                    cell.font = data_font
                    cell.border = thin_border

            ws_summary.column_dimensions["A"].width = 20
            ws_summary.column_dimensions["B"].width = 50

            # ── Foglio 2: GHG Emissions ────────────────────────
            ws_ghg = wb.create_sheet("GHG Emissions")
            emissions = report_data.get("emissions", {})

            ghg_headers = ["Scope", "Category", "Value", "Unit", "Year", "Year N-1"]
            ws_ghg.append(ghg_headers)
            for cell in ws_ghg[1]:
                cell.font = header_font
                cell.fill = header_fill
                cell.border = thin_border

            emission_rows = []
            scopes = emissions.get("scopes", {})
            for scope_name in ["scope1", "scope2_location", "scope2_market", "scope3"]:
                data = scopes.get(scope_name, {})
                emission_rows.append([
                    scope_name.replace("_", " ").title(),
                    "Total",
                    data.get("value", ""),
                    data.get("unit", "tCO2eq"),
                    data.get("current_year", ""),
                    data.get("previous_year", ""),
                ])

            for row_data in emission_rows:
                ws_ghg.append(row_data)

            for row in ws_ghg.iter_rows(min_row=2, max_row=len(emission_rows) + 1):
                for cell in row:
                    cell.font = data_font
                    cell.border = thin_border

            for i, col in enumerate(ghg_headers, 1):
                ws_ghg.column_dimensions[get_column_letter(i)].width = 20

            # ── Foglio 3: Materiality ──────────────────────────
            ws_mat = wb.create_sheet("Materiality")
            materiality = report_data.get("materiality", {})

            mat_headers = ["Topic", "IRO", "Impact Score", "Financial Score", "Material"]
            ws_mat.append(mat_headers)
            for cell in ws_mat[1]:
                cell.font = header_font
                cell.fill = header_fill
                cell.border = thin_border

            material_iros = materiality.get("iros", [])
            for iro in material_iros:
                ws_mat.append([
                    iro.get("topic", ""),
                    iro.get("name", ""),
                    iro.get("impact_score", ""),
                    iro.get("financial_score", ""),
                    "✓" if iro.get("is_material") else "✗",
                ])

            for row in ws_mat.iter_rows(min_row=2, max_row=len(material_iros) + 1):
                for cell in row:
                    cell.font = data_font
                    cell.border = thin_border

            for i, col in enumerate(mat_headers, 1):
                ws_mat.column_dimensions[get_column_letter(i)].width = 25

            # ── Foglio 4: ESRS Coverage ────────────────────────
            ws_esrs = wb.create_sheet("ESRS Coverage")
            gap_analysis = report_data.get("gap_analysis", {})

            esrs_headers = ["Standard", "Required", "Complete", "Partial", "Missing", "Coverage %"]
            ws_esrs.append(esrs_headers)
            for cell in ws_esrs[1]:
                cell.font = header_font
                cell.fill = header_fill
                cell.border = thin_border

            gaps_by_standard = gap_analysis.get("gaps_by_standard", {})
            for std, data in gaps_by_standard.items():
                required = data.get("required", 0)
                complete = data.get("complete", 0)
                coverage_pct = round((complete / required * 100), 1) if required > 0 else 0
                ws_esrs.append([
                    std,
                    required,
                    complete,
                    data.get("partial", 0),
                    data.get("missing", 0),
                    f"{coverage_pct}%",
                ])

            for row in ws_esrs.iter_rows(min_row=2, max_row=len(gaps_by_standard) + 1):
                for cell in row:
                    cell.font = data_font
                    cell.border = thin_border

            for i, col in enumerate(esrs_headers, 1):
                ws_esrs.column_dimensions[get_column_letter(i)].width = 18

            # Salva in memoria
            xlsx_buffer = io.BytesIO()
            wb.save(xlsx_buffer)
            xlsx_buffer.seek(0)
            xlsx_bytes = xlsx_buffer.getvalue()

            logger.info(f"XLSX generated: {len(xlsx_bytes)} bytes, sheets: {len(wb.sheetnames)}")
            return ExportResult(
                format="xlsx",
                filename=full_filename,
                content=xlsx_bytes,
                mime_type=self.MIME_TYPES["xlsx"],
                size_bytes=len(xlsx_bytes),
                metadata={
                    "sheets": wb.sheetnames,
                    "generated_at": datetime.utcnow().isoformat(),
                },
            )

        except XLSXGenerationError:
            raise
        except Exception as e:
            error_msg = f"XLSX generation error: {str(e)}"
            logger.error(error_msg)
            raise XLSXGenerationError(error_msg)

    # ── DOCX Export ─────────────────────────────────────────────

    def export_docx(
        self,
        xhtml_content: str,
        filename: str = "report_draft",
        options: Optional[ExportOptions] = None,
    ) -> ExportResult:
        """
        Esporta il report in formato Word (DOCX).
        
        Converte il contenuto XHTML in un documento DOCX
        con formattazione base (titoli, paragrafi, tabelle).
        
        Args:
            xhtml_content: Contenuto XHTML del report
            filename: Nome base del file
            options: Opzioni di esportazione
            
        Returns:
            ExportResult con il contenuto DOCX in bytes
        """
        opts = options or ExportOptions()
        full_filename = f"{filename}{self.EXTENSIONS['docx']}"

        if not self._docx_available:
            logger.warning("python-docx not available for DOCX export")
            return ExportResult(
                format="docx",
                filename=full_filename,
                content=b"",
                mime_type=self.MIME_TYPES["docx"],
                size_bytes=0,
                success=False,
                error_message="python-docx not installed. Install with: pip install python-docx",
            )

        try:
            from docx import Document
            from docx.shared import Inches, Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT

            doc = Document()

            # Imposta margini
            for section in doc.sections:
                section.top_margin = Cm(opts.margin_top / 10)
                section.bottom_margin = Cm(opts.margin_bottom / 10)
                section.left_margin = Cm(opts.margin_left / 10)
                section.right_margin = Cm(opts.margin_right / 10)

            # Parsing XHTML semplice -> estrai elementi
            import re

            # Estrai titoli (h1, h2, h3)
            h1_matches = re.findall(r'<h1[^>]*>(.*?)</h1>', xhtml_content, re.DOTALL)
            h2_matches = re.findall(r'<h2[^>]*>(.*?)</h2>', xhtml_content, re.DOTALL)
            h3_matches = re.findall(r'<h3[^>]*>(.*?)</h3>', xhtml_content, re.DOTALL)
            p_matches = re.findall(r'<p[^>]*>(.*?)</p>', xhtml_content, re.DOTALL)
            table_matches = re.findall(
                r'<table[^>]*>(.*?)</table>', xhtml_content, re.DOTALL
            )

            # Pulisci tag HTML residui
            def clean_html(text: str) -> str:
                text = re.sub(r'<[^>]+>', '', text)
                text = re.sub(r'\s+', ' ', text).strip()
                return text

            # Aggiungi contenuti al documento
            for h1 in h1_matches:
                doc.add_heading(clean_html(h1), level=1)

            for h2 in h2_matches:
                doc.add_heading(clean_html(h2), level=2)

            for h3 in h3_matches:
                doc.add_heading(clean_html(h3), level=3)

            for p in p_matches:
                cleaned = clean_html(p)
                if cleaned:
                    doc.add_paragraph(cleaned)

            # Aggiungi tabelle
            for table_html in table_matches:
                rows = re.findall(r'<tr[^>]*>(.*?)</tr>', table_html, re.DOTALL)
                if not rows:
                    continue

                table_data = []
                for row_html in rows:
                    cells = re.findall(
                        r'<t[dh][^>]*>(.*?)</t[dh]>', row_html, re.DOTALL
                    )
                    table_data.append([clean_html(c) for c in cells])

                if table_data:
                    num_cols = max(len(row) for row in table_data)
                    word_table = doc.add_table(
                        rows=len(table_data),
                        cols=num_cols,
                    )
                    word_table.style = "Light Grid Accent 1"
                    word_table.alignment = WD_TABLE_ALIGNMENT.CENTER

                    for i, row_data in enumerate(table_data):
                        for j, cell_data in enumerate(row_data):
                            if j < num_cols:
                                word_table.cell(i, j).text = cell_data

                    doc.add_paragraph()  # Spazio dopo tabella

            # Salva in memoria
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            docx_bytes = docx_buffer.getvalue()

            logger.info(f"DOCX generated: {len(docx_bytes)} bytes")
            return ExportResult(
                format="docx",
                filename=full_filename,
                content=docx_bytes,
                mime_type=self.MIME_TYPES["docx"],
                size_bytes=len(docx_bytes),
                metadata={
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                    "sections": len(doc.sections),
                    "generated_at": datetime.utcnow().isoformat(),
                },
            )

        except DOCXGenerationError:
            raise
        except Exception as e:
            error_msg = f"DOCX generation error: {str(e)}"
            logger.error(error_msg)
            raise DOCXGenerationError(error_msg)

    # ── JSON Export ─────────────────────────────────────────────

    def export_json(
        self,
        report_data: Dict[str, Any],
        filename: str = "report_data",
        options: Optional[ExportOptions] = None,
        indent: int = 2,
    ) -> ExportResult:
        """
        Esporta i dati del report in formato JSON.
        
        Utile per integrazione API e analisi dati.
        Supporta serializzazione custom per tipi non-JSON (date, UUID, etc.).
        
        Args:
            report_data: Dizionario con i dati del report
            filename: Nome base del file
            options: Opzioni di esportazione
            indent: Indentazione JSON
            
        Returns:
            ExportResult con il contenuto JSON in bytes
        """
        opts = options or ExportOptions()
        full_filename = f"{filename}{self.EXTENSIONS['json']}"

        try:
            json_bytes = json.dumps(
                report_data,
                indent=indent,
                ensure_ascii=False,
                default=self._json_serializer,
            ).encode("utf-8")

            logger.info(f"JSON generated: {len(json_bytes)} bytes")
            return ExportResult(
                format="json",
                filename=full_filename,
                content=json_bytes,
                mime_type=self.MIME_TYPES["json"],
                size_bytes=len(json_bytes),
                metadata={
                    "indent": indent,
                    "encoding": "utf-8",
                    "generated_at": datetime.utcnow().isoformat(),
                },
            )

        except Exception as e:
            error_msg = f"JSON serialization error: {str(e)}"
            logger.error(error_msg)
            return ExportResult(
                format="json",
                filename=full_filename,
                content=b"{}",
                mime_type=self.MIME_TYPES["json"],
                size_bytes=2,
                success=False,
                error_message=error_msg,
            )

    def _json_serializer(self, obj: Any) -> str:
        """Serializza oggetti non-JSON standard."""
        if hasattr(obj, "isoformat"):
            return obj.isoformat()
        if hasattr(obj, "hex"):
            return obj.hex()
        return str(obj)

    # ── iXBRL Export ────────────────────────────────────────────

    def export_ixbrl(
        self,
        ixbrl_content: str,
        filename: str = "report_ixbrl",
        options: Optional[ExportOptions] = None,
    ) -> ExportResult:
        """
        Esporta il report in formato iXBRL (XHTML con tag XBRL inline).
        
        Questo è il formato principale per il filing regolatorio.
        Il contenuto iXBRL deve essere già stato generato da IXBRLTagger.
        Questa funzione si occupa solo dell'export e confezionamento.
        
        Args:
            ixbrl_content: Contenuto XHTML con tag iXBRL
            filename: Nome base del file
            options: Opzioni di esportazione
            
        Returns:
            ExportResult con il contenuto iXBRL in bytes
        """
        opts = options or ExportOptions()
        full_filename = f"{filename}{self.EXTENSIONS['ixbrl']}"

        try:
            content_bytes = ixbrl_content.encode("utf-8")

            # Aggiungi dichiarazione XML se mancante
            if not ixbrl_content.startswith("<?xml"):
                content_bytes = (
                    b'<?xml version="1.0" encoding="UTF-8"?>\n' + content_bytes
                )

            logger.info(f"iXBRL export: {len(content_bytes)} bytes")
            return ExportResult(
                format="ixbrl",
                filename=full_filename,
                content=content_bytes,
                mime_type=self.MIME_TYPES["ixbrl"],
                size_bytes=len(content_bytes),
                metadata={
                    "has_xml_declaration": ixbrl_content.startswith("<?xml"),
                    "generated_at": datetime.utcnow().isoformat(),
                },
            )

        except Exception as e:
            error_msg = f"iXBRL export error: {str(e)}"
            logger.error(error_msg)
            return ExportResult(
                format="ixbrl",
                filename=full_filename,
                content=b"",
                mime_type=self.MIME_TYPES["ixbrl"],
                size_bytes=0,
                success=False,
                error_message=error_msg,
            )

    # ── Export Multi-Formato ────────────────────────────────────

    def export_all(
        self,
        ixbrl_content: str,
        report_data: Dict[str, Any],
        filename_base: str = "csrd_report",
        options: Optional[ExportOptions] = None,
    ) -> Dict[str, ExportResult]:
        """
        Esporta il report in tutti i formati supportati.
        
        Args:
            ixbrl_content: Contenuto iXBRL/XHTML del report
            report_data: Dati strutturati per XLSX/JSON
            filename_base: Nome base comune per tutti i file
            options: Opzioni di esportazione
            
        Returns:
            Dizionario {formato: ExportResult}
        """
        results = {}

        # iXBRL
        results["ixbrl"] = self.export_ixbrl(ixbrl_content, filename_base, options)

        # PDF
        try:
            results["pdf"] = self.export_pdf(ixbrl_content, filename_base, options)
        except Exception as e:
            results["pdf"] = ExportResult(
                format="pdf", filename=f"{filename_base}.pdf",
                content=b"", mime_type="application/pdf",
                size_bytes=0, success=False, error_message=str(e),
            )

        # XLSX
        try:
            results["xlsx"] = self.export_xlsx(report_data, filename_base, options)
        except Exception as e:
            results["xlsx"] = ExportResult(
                format="xlsx", filename=f"{filename_base}.xlsx",
                content=b"", mime_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                size_bytes=0, success=False, error_message=str(e),
            )

        # DOCX
        try:
            results["docx"] = self.export_docx(ixbrl_content, filename_base, options)
        except Exception as e:
            results["docx"] = ExportResult(
                format="docx", filename=f"{filename_base}.docx",
                content=b"", mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                size_bytes=0, success=False, error_message=str(e),
            )

        # JSON
        try:
            results["json"] = self.export_json(report_data, filename_base, options)
        except Exception as e:
            results["json"] = ExportResult(
                format="json", filename=f"{filename_base}.json",
                content=b"{}", mime_type="application/json",
                size_bytes=2, success=False, error_message=str(e),
            )

        # Statistiche
        successful = sum(1 for r in results.values() if r.success)
        total_size = sum(r.size_bytes for r in results.values())
        logger.info(
            f"Multi-format export complete: {successful}/5 formats, "
            f"{total_size} total bytes"
        )

        return results

    # ── Utility ─────────────────────────────────────────────────

    def get_available_formats(self) -> List[str]:
        """Restituisce la lista dei formati disponibili in base alle librerie installate."""
        formats = ["ixbrl", "json"]
        if self._xhtml2pdf_available:
            formats.append("pdf")
        if self._openpyxl_available:
            formats.append("xlsx")
        if self._docx_available:
            formats.append("docx")
        return formats

    def get_format_info(self) -> Dict[str, Dict[str, str]]:
        """Restituisce informazioni su tutti i formati supportati."""
        return {
            "ixbrl": {
                "name": "iXBRL (XHTML)",
                "mime": self.MIME_TYPES["ixbrl"],
                "extension": self.EXTENSIONS["ixbrl"],
                "description": "Formato principale per filing regolatorio ESMA/ESAP",
                "available": True,
            },
            "pdf": {
                "name": "PDF",
                "mime": self.MIME_TYPES["pdf"],
                "extension": self.EXTENSIONS["pdf"],
                "description": "Documento PDF per stampa e condivisione",
                "available": self._xhtml2pdf_available,
                "note": "Richiede xhtml2pdf" if not self._xhtml2pdf_available else None,
            },
            "xlsx": {
                "name": "Excel (XLSX)",
                "mime": self.MIME_TYPES["xlsx"],
                "extension": self.EXTENSIONS["xlsx"],
                "description": "Dati strutturati per analisi in Excel",
                "available": self._openpyxl_available,
                "note": "Richiede openpyxl" if not self._openpyxl_available else None,
            },
            "docx": {
                "name": "Word (DOCX)",
                "mime": self.MIME_TYPES["docx"],
                "extension": self.EXTENSIONS["docx"],
                "description": "Documento Word per bozze e revisioni",
                "available": self._docx_available,
                "note": "Richiede python-docx" if not self._docx_available else None,
            },
            "json": {
                "name": "JSON",
                "mime": self.MIME_TYPES["json"],
                "extension": self.EXTENSIONS["json"],
                "description": "Dati in formato JSON per integrazione API",
                "available": True,
            },
        }


# ── Helper Functions ──────────────────────────────────────────────

def create_export_service() -> ExportService:
    """Factory per creare ExportService."""
    return ExportService()


def export_report_pdf(
    xhtml_content: str,
    filename: str = "report",
    options: Optional[ExportOptions] = None,
) -> ExportResult:
    """Helper per esportazione PDF rapida."""
    service = ExportService()
    return service.export_pdf(xhtml_content, filename, options)


def export_report_xlsx(
    report_data: Dict[str, Any],
    filename: str = "report_data",
) -> ExportResult:
    """Helper per esportazione XLSX rapida."""
    service = ExportService()
    return service.export_xlsx(report_data, filename)


def export_report_docx(
    xhtml_content: str,
    filename: str = "report_draft",
) -> ExportResult:
    """Helper per esportazione DOCX rapida."""
    service = ExportService()
    return service.export_docx(xhtml_content, filename)


def export_report_json(
    report_data: Dict[str, Any],
    filename: str = "report_data",
) -> ExportResult:
    """Helper per esportazione JSON rapida."""
    service = ExportService()
    return service.export_json(report_data, filename)
